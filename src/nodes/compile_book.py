import io
import logging

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from src.deps import get_supabase
from src.state import BookState

logger = logging.getLogger(__name__)


def compile_book(state: BookState) -> BookState:
    """
    Compile all approved chapters into a DOCX and TXT,
    upload DOCX to Supabase Storage and update book record.
    """
    supabase = get_supabase()

    # Reload final review info from DB for safety
    book_row = (
        supabase.table("books")
        .select("final_review_notes_status, final_review_notes, title")
        .eq("id", state.book_id)
        .single()
        .execute()
        .data
    )

    fr_status = book_row.get("final_review_notes_status")
    fr_notes = book_row.get("final_review_notes") or ""

    if not (
        fr_status == "no_notes_needed" or (fr_notes and fr_notes.strip())
    ):
        logger.warning("Final review not ready - missing notes or status")
        state.control["event"] = "final_review_not_ready"
        return state

    logger.info("Compiling final book...")
    # Fetch chapters in order
    resp = (
        supabase.table("chapters")
        .select("*")
        .eq("book_id", state.book_id)
        # supabase-py versions differ; default is ascending if not specified
        .order("chapter_number")
        .execute()
    )
    chapters = resp.data or []
    logger.info(f"Found {len(chapters)} chapters to compile")

    # Build text content
    txt_content = f"{book_row.get('title') or state.title}\n\n"
    txt_content += "\n\n".join(
        f"{c['title']}\n\n{c['content'] or ''}" for c in chapters
    )

    # Upload TXT to Supabase Storage
    bucket = "books-output"
    txt_filename = f"{state.book_id}.txt"
    logger.info(f"Uploading {txt_filename} to Supabase Storage...")
    storage = supabase.storage.from_(bucket)
    storage.upload(txt_filename, txt_content.encode("utf-8"))
    txt_url = storage.get_public_url(txt_filename)
    logger.info(f"TXT upload complete: {txt_url}")

    # Build DOCX
    doc = Document()
    doc.add_heading(book_row.get("title") or state.title, level=0)
    for c in chapters:
        doc.add_heading(c["title"], level=1)
        content = c["content"] or ""
        for para in content.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    # Upload DOCX to Supabase Storage
    docx_filename = f"{state.book_id}.docx"
    logger.info(f"Uploading {docx_filename} to Supabase Storage...")
    storage.upload(docx_filename, buf.getvalue())
    docx_url = storage.get_public_url(docx_filename)
    logger.info(f"DOCX upload complete: {docx_url}")

    # Build PDF
    pdf_buf = io.BytesIO()
    pdf_doc = SimpleDocTemplate(pdf_buf, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    # Title
    title_text = book_row.get("title") or state.title
    story.append(Paragraph(title_text, styles["Title"]))
    story.append(Spacer(1, 0.5 * inch))

    # Chapters
    for c in chapters:
        story.append(Paragraph(c["title"], styles["Heading1"]))
        content = c["content"] or ""
        for para in content.split("\n\n"):
            if para.strip():
                # Simple text cleaning for PDF
                para_clean = para.strip().replace("\n", "<br/>")
                story.append(Paragraph(para_clean, styles["Normal"]))
                story.append(Spacer(1, 0.2 * inch))

    try:
        pdf_doc.build(story)
        pdf_buf.seek(0)

        # Upload PDF to Supabase Storage
        pdf_filename = f"{state.book_id}.pdf"
        logger.info(f"Uploading {pdf_filename} to Supabase Storage...")
        storage.upload(pdf_filename, pdf_buf.getvalue())
        pdf_url = storage.get_public_url(pdf_filename)
        logger.info(f"PDF upload complete: {pdf_url}")
    except Exception as e:
        logger.warning(f"PDF generation failed: {e} - continuing without PDF")
        pdf_url = None

    # Store DOCX URL as primary output
    supabase.table("books").update(
        {
            "book_output_status": "ready",
            "output_file_url": docx_url,
        }
    ).eq("id", state.book_id).execute()

    state.book_output_status = "ready"
    state.control["event"] = "book_compiled"
    logger.info("Book compilation complete!")
    return state

